import pandas as pd
import matplotlib.pyplot as plt
from collections import defaultdict as ddict
import os
from docx import Document
from docx.shared import Inches
# Make sure working from corect directery
try:
    os.chdir(r'C:\Users\peter\Desktop\Python Projects\Hildas Day')
    print(os.getcwd())
except OSError:
    print('Current working directery')
    pass

# Create the word document:
document = Document()

# accessing the file using the absolute path
file_name = 'hildas_day.csv'

# get the csv data into a pandas data frame
data = pd.read_csv(file_name, encoding='Latin-1')

# total number of responses just grab length of any of the columns
print(f"Total responds numbers = {len(data['Timestamp'])}")
document.add_heading('Hildas Day', 0)
document.add_paragraph('Total number of responses:' + str(len(data['Timestamp'])))

def analyse_data(data_title, header, nan='Other'):
      ''' Function analyses the data and outputs a plot '''
      print('analysing ' + data_title)
      data_name = data[header]
      data_sizes = ddict(int)
      for data_point in data_name:
          data_sizes[data_point] += 1

      # Extract Data and titles
      data_wedges = data_sizes.values()
      data_titles = data_sizes.keys()

      # Plot the data
      plt.figure()
      plt.title(data_title)
      chart, texts, gr = plt.pie(data_wedges,colors=None, autopct = '%1.1f%%')
      plt.legend(chart, data_titles, loc='upper right', bbox_to_anchor=(1, 1), bbox_transform=plt.gcf().transFigure, fontsize='6')
      plt.savefig(data_title)

      # Add the charts to the word document
      document.add_picture(data_title + '.png', width=Inches(5))
      return

# Generate all the charts
analyse_data('Year Level', "What year are you?")
analyse_data('Gender', 'Which of the following best describes you?')
analyse_data('Activity', "What morning activity will you like to participate in?")
analyse_data('Laundry Votes', "Following the early morning event it has been common to then go to Laundry Bar for a morning rave. Do you like this idea?")
analyse_data('Laundry Alternatives', "If you are unhappy with laundry?")

# Finally add all the 'other' responses to the document in a list
document.add_heading('Other Repsonses:', 1)
for line in data["If you have any questions or issues please fill this question out, change can't happen unless someone speaks up. "]:
    document.add_paragraph(str(line), style='List Bullet')
document.save('Hildas Day Analysed.docx')
